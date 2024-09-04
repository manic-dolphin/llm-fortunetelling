import os
from tqdm import tqdm
import pandas as pd
from docx import Document

# doc = Document("data\\raw_data\\word\\2023-08-07_铁口箓国学文化八字实战擂台赛第364期反馈.docx")

def read_docx(file_path):
    # 读取.docx文件
    doc = Document(file_path)
    content = " ".join([para.text.strip() for para in doc.paragraphs])
    return content

def main(dictionary):
    data = []
    # 遍历目录下所有.docx文件
    for root, dirs, files in os.walk(dictionary):
        for file in tqdm(files):
            if file.endswith('.docx'):
                file_path = os.path.join(root, file)
                content = read_docx(file_path)
                data.append(content)
    
    df = pd.DataFrame(data)
    df.to_csv("output.csv", index=False, encoding='utf-8')

    return df

if __name__ == "__main__":
    df = main("data\\raw_data")
    df = pd.read_csv("output.csv")
    print(df.head())